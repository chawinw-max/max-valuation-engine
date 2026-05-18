"""
Consolidation Report generator — produces a .docx audit report
for the IB team to review the consolidation work.
"""

import io
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ── Helpers ──────────────────────────────────────────────────────────────────

def _fmt_thb(value):
    """Format a number as Thai Baht."""
    if value is None:
        return "N/A"
    try:
        v = float(value)
        if abs(v) >= 1_000_000:
            return f"THB {v / 1_000_000:,.1f}M"
        return f"THB {v:,.0f}"
    except (ValueError, TypeError):
        return str(value)


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for ci, header in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val) if val is not None else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)

    return table


def _calc_derived_from_pnl(key, pnl, year):
    """Calculate derived P&L totals for report tables."""
    def _v(k):
        return (pnl.get(k) or {}).get(year) or 0

    if key == "revenue":
        return _v("sales_and_services") + _v("other_revenues")
    elif key == "gross_profit":
        rev = _v("sales_and_services") + _v("other_revenues")
        return rev - abs(_v("cost_of_goods_sold"))
    elif key == "ebitda":
        gp = _calc_derived_from_pnl("gross_profit", pnl, year)
        return gp - abs(_v("selling_expenses")) - abs(_v("administrative_expenses")) - abs(_v("other_expenses"))
    elif key == "ebit":
        ebitda = _calc_derived_from_pnl("ebitda", pnl, year)
        return ebitda - abs(_v("depreciation_amortization"))
    elif key == "pbt":
        ebit = _calc_derived_from_pnl("ebit", pnl, year)
        return ebit - abs(_v("interest_expense"))
    elif key == "net_profit":
        pbt = _calc_derived_from_pnl("pbt", pnl, year)
        return pbt - abs(_v("income_tax"))
    return 0


# ── Main generator ───────────────────────────────────────────────────────────

def generate_consolidation_report(
    extractions: list,
    adjustments: list,
    intercompany: list,
    narrative: dict,
    entity_configs: list,
    fiscal_years: list,
    deal_code: str = "",
    source_files: list = None,
) -> bytes:
    """Generate the Word consolidation report and return as bytes."""

    doc = Document()
    years = [str(y) for y in sorted(fiscal_years)]
    entities = [e["name"] for e in entity_configs]

    # ── Page setup ───────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading("Financial Statement Consolidation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_info = [
        f"Deal: {deal_code}" if deal_code else None,
        f"Entities: {', '.join(entities)}",
        f"Fiscal Years: {', '.join(years)}",
        f"Prepared by MAX Solutions | {date.today().strftime('%B %d, %Y')}",
    ]
    for line in subtitle_info:
        if line:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ── Section 1: Executive Summary ─────────────────────────────────────────
    _add_heading(doc, "1. Executive Summary")
    doc.add_paragraph(narrative.get("executive_summary", "No executive summary generated."))

    # ── Section 2: Source Documents ──────────────────────────────────────────
    _add_heading(doc, "2. Source Documents")
    if source_files:
        headers = ["Filename", "Type"]
        rows = []
        for sf in source_files:
            name = sf if isinstance(sf, str) else sf.get("name", "Unknown")
            ftype = "Meeting Notes" if "note" in name.lower() or "meeting" in name.lower() else "Financial Document"
            rows.append([name, ftype])
        _add_table(doc, headers, rows, col_widths=[10, 5])
    else:
        doc.add_paragraph("Source document list not available.")

    doc.add_paragraph("")  # spacer

    if narrative.get("data_quality_notes"):
        p = doc.add_paragraph()
        p.add_run("Data Quality Notes: ").bold = True
        p.add_run(narrative["data_quality_notes"])

    # ── Section 3: Extracted Financials (As Reported) ────────────────────────
    _add_heading(doc, "3. Extracted Financials (As Reported)")

    _full_pnl_lines = [
        ("sales_and_services", "Revenue", False),
        ("other_revenues", "Other Revenue", False),
        (None, "gross_profit", True, "Gross Profit"),
        ("cost_of_goods_sold", "COGS", False),
        ("selling_expenses", "Selling Expenses", False),
        ("administrative_expenses", "Admin Expenses (excl. D&A)", False),
        ("other_expenses", "Other Expenses", False),
        (None, "ebitda", True, "EBITDA"),
        ("depreciation_amortization", "Depreciation & Amortization", False),
        (None, "ebit", True, "EBIT"),
        ("interest_expense", "Interest Expense", False),
        (None, "pbt", True, "Profit Before Tax"),
        ("income_tax", "Income Tax", False),
        (None, "net_profit", True, "Net Profit"),
    ]

    def _build_pnl_rows(pnl):
        rows = []
        for item in _full_pnl_lines:
            if item[2]:  # derived row
                label = item[3]
                calc_key = item[1]
                row = [label]
                for yr in years:
                    row.append(_fmt_thb(_calc_derived_from_pnl(calc_key, pnl, yr)))
                rows.append(row)
            else:
                key, label = item[0], item[1]
                row = [label]
                for yr in years:
                    val = (pnl.get(key) or {}).get(yr)
                    row.append(_fmt_thb(val))
                rows.append(row)
        return rows

    for ext in extractions:
        ename = ext.get("entity_name", "Unknown")
        pnl = ext.get("pnl", {})

        doc.add_heading(ename, level=2)
        headers = ["Line Item"] + years
        _add_table(doc, headers, _build_pnl_rows(pnl), col_widths=[5] + [3.5] * len(years))

        if ext.get("ai_notes"):
            doc.add_paragraph("")
            p = doc.add_paragraph()
            p.add_run("AI Notes: ").bold = True
            p.add_run(ext["ai_notes"])

    # Consolidated P&L (multi-entity)
    if len(extractions) > 1:
        doc.add_heading("Consolidated P&L", level=2)
        doc.add_paragraph("Sum of all entities for each fiscal year.")

        # Build consolidated PNL
        consol_pnl = {}
        for ext in extractions:
            pnl = ext.get("pnl", {})
            for k, v in pnl.items():
                if k not in consol_pnl:
                    consol_pnl[k] = {yr: 0 for yr in years}
                for yr in years:
                    consol_pnl[k][yr] = (consol_pnl[k].get(yr) or 0) + ((v or {}).get(yr) or 0)

        headers = ["Line Item"] + years
        _add_table(doc, headers, _build_pnl_rows(consol_pnl), col_widths=[5] + [3.5] * len(years))

    # ── Section 4: Normalization Adjustments ─────────────────────────────────
    _add_heading(doc, "4. Normalization Adjustments")

    accepted = [a for a in adjustments if a.get("action") != "reject"]
    rejected = [a for a in adjustments if a.get("action") == "reject"]

    cat_labels = {
        "owner_expense": "Owner Personal Expenses",
        "related_party": "Related-Party Adjustments",
        "tax_adjustment": "Tax-Motivated Adjustments",
        "non_recurring": "Non-Recurring Items",
        "intercompany": "Intercompany Eliminations",
    }

    cat_narratives = narrative.get("adjustment_narratives", {})

    for cat, label in cat_labels.items():
        cat_adjs = [a for a in accepted if a.get("category") == cat]
        if not cat_adjs:
            continue

        doc.add_heading(label, level=2)

        # Narrative paragraph
        if cat_narratives.get(cat):
            doc.add_paragraph(cat_narratives[cat])

        # Adjustment detail table
        headers = ["Entity", "Year", "Line Item", "Original", "Adjustment", "EBITDA Impact", "Description"]
        rows = []
        for adj in cat_adjs:
            rows.append([
                adj.get("entity", ""),
                adj.get("year", ""),
                adj.get("line_item", "").replace("_", " ").title(),
                _fmt_thb(adj.get("original_amount")),
                _fmt_thb(adj.get("adjustment_amount")),
                _fmt_thb(adj.get("ebitda_impact")),
                adj.get("description", ""),
            ])
        _add_table(doc, headers, rows, col_widths=[3, 2, 3, 2.5, 2.5, 2.5, 4])

        # Detailed explanations
        for adj in cat_adjs:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            p.add_run(f"{adj.get('description', 'Adjustment')}: ").bold = True
            p.add_run(adj.get("explanation", ""))
            evidence = adj.get("evidence_source", "")
            if evidence:
                p.add_run(f" [Source: {evidence}]").font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Rejected adjustments
    if rejected:
        doc.add_heading("Rejected Adjustments", level=2)
        doc.add_paragraph("The following proposed adjustments were reviewed and rejected by the analyst:")
        for adj in rejected:
            doc.add_paragraph(
                f"- {adj.get('description', 'Adjustment')} ({adj.get('entity', '')}, "
                f"{adj.get('year', '')}): {_fmt_thb(adj.get('adjustment_amount'))} — "
                f"{adj.get('explanation', '')}",
            )

    # ── Section 5: Intercompany Eliminations ─────────────────────────────────
    if intercompany:
        _add_heading(doc, "5. Intercompany Eliminations")
        headers = ["From", "To", "Year", "Amount", "Description"]
        rows = []
        for ic in intercompany:
            if ic.get("action") == "reject":
                continue
            rows.append([
                ic.get("from_entity", ""),
                ic.get("to_entity", ""),
                ic.get("year", ""),
                _fmt_thb(ic.get("amount")),
                ic.get("description", ""),
            ])
        if rows:
            _add_table(doc, headers, rows, col_widths=[3.5, 3.5, 2, 3, 5])
        else:
            doc.add_paragraph("No intercompany eliminations applied.")

    # ── Section 6: Normalized Financial Summary ──────────────────────────────
    section_num = 6 if intercompany else 5
    _add_heading(doc, f"{section_num}. Normalized Financial Summary")

    # EBITDA Bridge table
    doc.add_heading("EBITDA Bridge", level=2)

    from core.consolidation_excel import _apply_adjustments, _calc_derived

    def _consol_pnl(kind):
        combined = {}
        for ext in extractions:
            ename = ext.get("entity_name", "")
            pnl = ext.get("pnl", {})
            if kind == "normalized":
                pnl = _apply_adjustments(pnl, accepted, ename, years)
            for k, v in pnl.items():
                if k not in combined:
                    combined[k] = {yr: 0 for yr in years}
                for yr in years:
                    combined[k][yr] = (combined[k].get(yr) or 0) + ((v or {}).get(yr) or 0)
        return combined

    reported_pnl = _consol_pnl("reported")
    normalized_pnl = _consol_pnl("normalized")

    headers = [""] + years
    bridge_rows = [["Reported EBITDA"] + [_fmt_thb(_calc_derived("EBITDA", reported_pnl, yr)) for yr in years]]

    for cat, label in cat_labels.items():
        cat_adjs = [a for a in accepted if a.get("category") == cat]
        if not cat_adjs:
            continue
        row = [f"  + {label}"]
        for yr in years:
            total = sum(a.get("ebitda_impact", 0) or 0 for a in cat_adjs if a.get("year") == yr)
            row.append(_fmt_thb(total) if total else "-")
        bridge_rows.append(row)

    bridge_rows.append(
        ["Normalized EBITDA"] + [_fmt_thb(_calc_derived("EBITDA", normalized_pnl, yr)) for yr in years]
    )

    _add_table(doc, headers, bridge_rows, col_widths=[5] + [3] * len(years))

    # ── Section 7: Flags & Warnings ──────────────────────────────────────────
    section_num += 1
    _add_heading(doc, f"{section_num}. Flags & Warnings for IB Team")

    if narrative.get("warnings_section"):
        doc.add_paragraph(narrative["warnings_section"])
    else:
        doc.add_paragraph("No additional flags or warnings.")

    # ── Section 8: Meeting Notes Summary ─────────────────────────────────────
    section_num += 1
    _add_heading(doc, f"{section_num}. Meeting Notes Summary")

    if narrative.get("meeting_notes_section"):
        doc.add_paragraph(narrative["meeting_notes_section"])
    else:
        doc.add_paragraph("No meeting notes were provided for this consolidation.")

    # ── Section 9: Conclusion ────────────────────────────────────────────────
    section_num += 1
    _add_heading(doc, f"{section_num}. Conclusion")

    if narrative.get("conclusion"):
        doc.add_paragraph(narrative["conclusion"])

    # ── Disclaimer ───────────────────────────────────────────────────────────
    doc.add_page_break()
    disclaimer = doc.add_paragraph()
    disclaimer.add_run("Disclaimer: ").bold = True
    disclaimer.add_run(
        "This consolidation report was prepared using AI-assisted analysis and is intended "
        "as a working document for the MAX Solutions investment banking team. All adjustments "
        "and figures should be independently verified before use in any client-facing materials, "
        "valuation models, or transaction documents. MAX Solutions accepts no liability for "
        "errors arising from AI-generated analysis or incomplete source data."
    )
    disclaimer.runs[-1].font.size = Pt(9)
    disclaimer.runs[-1].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Save to bytes ────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
